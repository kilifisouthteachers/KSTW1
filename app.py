import sqlite3
import random
import string
import csv
import tkinter as tk
import csv
import smtplib
import stripe
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from fpdf import FPDF
from tkinter import messagebox
from docx import Document
conn = sqlite3.connect('kstw_database.db')
cursor = conn.cursor()

cursor.execute('''CREATE TABLE IF NOT EXISTS members (
                membership_number TEXT PRIMARY KEY,
                name TEXT,
                institution TEXT,
                cluster TEXT,
                safaricom_number TEXT,
                amount_contributed REAL DEFAULT 0)''')
conn.commit()

def display_menu():
    print("\nKSTW Member Registration App")
    print("1. Register Member")
    print("2. Make Contribution")
    print("3. Export Report")
    print("4. Exit")
    
def generate_membership_number():
    cursor.execute("SELECT COUNT(*) FROM members")
    count = cursor.fetchone()[0] + 1
    return f"KSTW{count:04d}-24"

def register_member():
    print("\nMember Registration")
    name = input("Enter full name: ")
    institution = input("Enter institution: ")
    cluster = input("Enter cluster: ")
    safaricom_number = input("Enter Safaricom number: ")
    membership_number = generate_membership_number()

    cursor.execute('''INSERT INTO members 
                    (membership_number, name, institution, cluster, safaricom_number)
                    VALUES (?, ?, ?, ?, ?)''',
                    (membership_number, name, institution, cluster, safaricom_number))
    conn.commit()
    print("Registration successful. Your membership number is:", membership_number)

def make_contribution():
    print("\nContribution")
    membership_number = input("Enter membership number: ")
    amount = float(input("Enter amount contributed: "))
    cursor.execute("UPDATE members SET amount_contributed = amount_contributed + ? WHERE membership_number = ?", (amount, membership_number))
    conn.commit()
    print("Contribution recorded successfully.")
    
def export_report(format_type):
    cursor.execute("SELECT * FROM members")
    data = cursor.fetchall()

    if format_type == 'csv':
        with open('kstw_report.csv', 'w', newline='') as f:
            writer = csv.writer(f)
            writer.writerow(['Membership Number', 'Name', 'Institution', 'Cluster', 'Safaricom Number', 'Amount Contributed'])
            writer.writerows(data)
        print("CSV report generated successfully.")
    elif format_type == 'docx':
        doc = Document()
        doc.add_heading('KSTW Member Report', level=1)
        for member in data:
            doc.add_paragraph(f"Membership Number: {member[0]}")
            doc.add_paragraph(f"Name: {member[1]}")
            doc.add_paragraph(f"Institution: {member[2]}")
            doc.add_paragraph(f"Cluster: {member[3]}")
            doc.add_paragraph(f"Safaricom Number: {member[4]}")
            doc.add_paragraph(f"Amount Contributed: {member[5]}")
            doc.add_page_break()
        doc.save('kstw_report.docx')
        print("DOCX report generated successfully.")
        
def main():
    while True:
        display_menu()
        choice = input("Enter your choice: ")

        if choice == '1':
            register_member()
        elif choice == '2':
            make_contribution()
        elif choice == '3':
            format_type = input("Enter format for report (csv/docx): ")
            export_report(format_type)
        elif choice == '4':
            break
        else:
            print("Invalid choice. Please try again.")

    conn.close()

if __name__ == "__main__":
    main()
    
conn = sqlite3.connect('user_database.db')
cursor = conn.cursor()

cursor.execute('''CREATE TABLE IF NOT EXISTS users (
                username TEXT PRIMARY KEY,
                password TEXT)''')
conn.commit()

def login():
    username = input("Enter username: ")
    password = input("Enter password: ")

    cursor.execute("SELECT * FROM users WHERE username = ? AND password = ?", (username, password))
    user = cursor.fetchone()

    if user:
        print("Login successful!")
    else:
        print("Invalid username or password.")

def register():
    username = input("Enter username: ")
    password = input("Enter password: ")

    cursor.execute("INSERT INTO users (username, password) VALUES (?, ?)", (username, password))
    conn.commit()
    print("Registration successful!")

def display_menu():
    print("\nUser Authentication")
    print("1. Login")
    print("2. Register")
    print("3. Exit")

def main():
    while True:
        display_menu()
        choice = input("Enter your choice: ")

        if choice == '1':
            login()
        elif choice == '2':
            register()
        elif choice == '3':
            break
        else:
            print("Invalid choice. Please try again.")

    conn.close()

if __name__ == "__main__":
    main()
    
conn = sqlite3.connect('kstw_database.db')
cursor = conn.cursor()

cursor.execute('''CREATE TABLE IF NOT EXISTS members (
                membership_number TEXT PRIMARY KEY,
                name TEXT,
                institution TEXT,
                cluster TEXT,
                safaricom_number TEXT,
                amount_contributed REAL DEFAULT 0)''')
conn.commit()

def register_member():
    membership_number = generate_membership_number()
    name = name_entry.get()
    institution = institution_entry.get()
    cluster = cluster_entry.get()
    safaricom_number = safaricom_entry.get()

    cursor.execute('''INSERT INTO members 
                    (membership_number, name, institution, cluster, safaricom_number)
                    VALUES (?, ?, ?, ?, ?)''',
                    (membership_number, name, institution, cluster, safaricom_number))
    conn.commit()
    messagebox.showinfo("Success", f"Registration successful! Your membership number is: {membership_number}")

def generate_membership_number():
    cursor.execute("SELECT COUNT(*) FROM members")
    count = cursor.fetchone()[0] + 1
    return f"KSTW{count:04d}-24"

root = tk.Tk()
root.title("KSTW Member Registration")

tk.Label(root, text="Name:").grid(row=0, column=0, sticky="w")
name_entry = tk.Entry(root)
name_entry.grid(row=0, column=1)

tk.Label(root, text="Institution:").grid(row=1, column=0, sticky="w")
institution_entry = tk.Entry(root)
institution_entry.grid(row=1, column=1)

tk.Label(root, text="Cluster:").grid(row=2, column=0, sticky="w")
cluster_entry = tk.Entry(root)
cluster_entry.grid(row=2, column=1)

tk.Label(root, text="Safaricom Number:").grid(row=3, column=0, sticky="w")
safaricom_entry = tk.Entry(root)
safaricom_entry.grid(row=3, column=1)

register_button = tk.Button(root, text="Register", command=register_member)
register_button.grid(row=4, column=1)

root.mainloop()
conn.close()

conn = sqlite3.connect('kstw_database.db')
cursor = conn.cursor()

cursor.execute('''CREATE TABLE IF NOT EXISTS members (
                membership_number TEXT PRIMARY KEY,
                name TEXT,
                institution TEXT,
                cluster TEXT,
                safaricom_number TEXT,
                amount_contributed REAL DEFAULT 0)''')
conn.commit()

def register_member():
    try:
        membership_number = generate_membership_number()
        name = name_entry.get()
        institution = institution_entry.get()
        cluster = cluster_entry.get()
        safaricom_number = safaricom_entry.get()

        cursor.execute('''INSERT INTO members 
                        (membership_number, name, institution, cluster, safaricom_number)
                        VALUES (?, ?, ?, ?, ?)''',
                        (membership_number, name, institution, cluster, safaricom_number))
        conn.commit()
        messagebox.showinfo("Success", f"Registration successful! Your membership number is: {membership_number}")
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {str(e)}")
def generate_membership_number():
    cursor.execute("SELECT COUNT(*) FROM members")
    count = cursor.fetchone()[0] + 1
    return f"KSTW{count:04d}-24"

root = tk.Tk()
root.title("KSTW Member Registration")

tk.Label(root, text="Name:").grid(row=0, column=0, sticky="w")
name_entry = tk.Entry(root)
name_entry.grid(row=0, column=1)

tk.Label(root, text="Institution:").grid(row=1, column=0, sticky="w")
institution_entry = tk.Entry(root)
institution_entry.grid(row=1, column=1)

tk.Label(root, text="Cluster:").grid(row=2, column=0, sticky="w")
cluster_entry = tk.Entry(root)
cluster_entry.grid(row=2, column=1)

tk.Label(root, text="Safaricom Number:").grid(row=3, column=0, sticky="w")
safaricom_entry = tk.Entry(root)
safaricom_entry.grid(row=3, column=1)

register_button = tk.Button(root, text="Register", command=register_member)
register_button.grid(row=4, column=1)
root.mainloop()
conn.close()

conn = sqlite3.connect('kstw_database.db')
cursor = conn.cursor()

cursor.execute('''CREATE TABLE IF NOT EXISTS members (
                membership_number TEXT PRIMARY KEY,
                name TEXT,
                institution TEXT,
                cluster TEXT,
                safaricom_number TEXT,
                amount_contributed REAL DEFAULT 0)''')
conn.commit()

def register_member():
    name = name_entry.get().strip()
    institution = institution_entry.get().strip()
    cluster = cluster_entry.get().strip()
    safaricom_number = safaricom_entry.get().strip()
    if not name or not institution or not cluster or not safaricom_number:
        messagebox.showerror("Error", "All fields are required.")
        return
    elif not safaricom_number.isdigit() or len(safaricom_number) != 10:
        messagebox.showerror("Error", "Safaricom number must be a 10-digit number.")
        return
    try:
        membership_number = generate_membership_number()
        cursor.execute('''INSERT INTO members 
                        (membership_number, name, institution, cluster, safaricom_number)
                        VALUES (?, ?, ?, ?, ?)''',
                        (membership_number, name, institution, cluster, safaricom_number))
        conn.commit()
        messagebox.showinfo("Success", f"Registration successful! Your membership number is: {membership_number}")
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {str(e)}")
def generate_membership_number():
    cursor.execute("SELECT COUNT(*) FROM members")
    count = cursor.fetchone()[0] + 1
    return f"KSTW{count:04d}-24"
root = tk.Tk()
root.title("KSTW Member Registration")

tk.Label(root, text="Name:").grid(row=0, column=0, sticky="w")
name_entry = tk.Entry(root)
name_entry.grid(row=0, column=1)

tk.Label(root, text="Institution:").grid(row=1, column=0, sticky="w")
institution_entry = tk.Entry(root)
institution_entry.grid(row=1, column=1)

tk.Label(root, text="Cluster:").grid(row=2, column=0, sticky="w")
cluster_entry = tk.Entry(root)
cluster_entry.grid(row=2, column=1)

tk.Label(root, text="Safaricom Number:").grid(row=3, column=0, sticky="w")
safaricom_entry = tk.Entry(root)
safaricom_entry.grid(row=3, column=1)

register_button = tk.Button(root, text="Register", command=register_member)
register_button.grid(row=4, column=1)
root.mainloop()
conn.close()

conn = sqlite3.connect('kstw_database.db')
cursor = conn.cursor()
def export_csv_report():
    cursor.execute("SELECT * FROM members")
    data = cursor.fetchall()

    with open('kstw_report.csv', 'w', newline='') as f:
        writer = csv.writer(f)
        writer.writerow(['Membership Number', 'Name', 'Institution', 'Cluster', 'Safaricom Number', 'Amount Contributed'])
        writer.writerows(data)

    print("CSV report generated successfully.")
def export_pdf_report():
    cursor.execute("SELECT * FROM members")
    data = cursor.fetchall()

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    pdf.cell(200, 10, txt="KSTW Member Report", ln=True, align="C")
    pdf.ln(10)

    # Add table headers
    pdf.set_font("Arial", size=10)
    pdf.cell(40, 10, "Membership Number", 1, 0, "C")
    pdf.cell(40, 10, "Name", 1, 0, "C")
    pdf.cell(40, 10, "Institution", 1, 0, "C")
    pdf.cell(40, 10, "Cluster", 1, 0, "C")
    pdf.cell(40, 10, "Safaricom Number", 1, 0, "C")
    pdf.cell(40, 10, "Amount Contributed", 1, 1, "C")
    pdf.set_font("Arial", size=8)
    for row in data:
        for col in row:
            pdf.cell(40, 10, str(col), 1, 0, "C")
        pdf.ln()

    pdf.output("kstw_report.pdf")

    print("PDF report generated successfully.")
while True:
    print("\nSelect report format:")
    print("1. Export CSV report")
    print("2. Export PDF report")
    print("3. Exit")

    choice = input("Enter your choice: ")

    if choice == '1':
        export_csv_report()
    elif choice == '2':
        export_pdf_report()
    elif choice == '3':
        break
    else:
        print("Invalid choice. Please try again.")

conn.close()

conn = sqlite3.connect('kstw_database.db')
cursor = conn.cursor().conn = sqlite3.connect('kstw_database.db')
cursor = conn.cursor()

cursor.execute("CREATE INDEX IF NOT EXISTS membership_number_index ON members (membership_number)")
conn.commit()
conn.close()

conn = sqlite3.connect('kstw_database.db')
cursor = conn.cursor()

def send_email_notification(email, membership_number):
    sender_email = "your_email@gmail.com"  # Replace with your email
    sender_password = "your_password"       # Replace with your email password

    subject = "KSTW Member Registration Confirmation"
    body = f"Dear Member,\n\nYour registration with membership number {membership_number} is successful.\n\nThank you for joining KSTW."

    message = MIMEMultipart()
    message["From"] = sender_email
    message["To"] = email
    message["Subject"] = subject

    message.attach(MIMEText(body, "plain"))
    
    try:
        smtp_server = smtplib.SMTP("smtp.gmail.com", 587)
        smtp_server.starttls()
        smtp_server.login(sender_email, sender_password)
        smtp_server.sendmail(sender_email, email, message.as_string())
        smtp_server.quit()
        print("Email notification sent successfully.")
    except Exception as e:
        print(f"Failed to send email notification: {e}")
def register_member(email):
    try:
        membership_number = "KSTW0001-24"
        cursor.execute('''INSERT INTO members 
                        (membership_number, name, institution, cluster, safaricom_number)
                        VALUES (?, ?, ?, ?, ?)''',
                        (membership_number, "John Doe", "Institution", "Cluster", "1234567890"))
        conn.commit()
        send_email_notification(email, membership_number)
    except Exception as e:
        print(f"Failed to register member: {e}")
email = "recipient_kazungukenny.com"
register_member(email)
conn.close()

conn = sqlite3.connect('kstw_database.db')
cursor = conn.cursor()
stripe.api_key = "s6gPrYcABSY3Oy1IDkzGytbaqEOg5fNP39Sv8uCTGYd3ZUu8jA9C2ZsCnjWBbwB1W0F0XrvgQHXhWGxDAsX31X48OrJaUWGSAFtj7Wc4zrdKtqY6blSw3p8jp4WRd7Yq"

def process_payment(amount, email):
    try:
        intent = stripe.PaymentIntent.create(
            amount=amount,
            currency='ksh',
            receipt_email=email
        )
        return intent.client_secret
    except stripe.error.StripeError as e:
        print(f"Failed to process payment: {e}")
        return None
def register_member_and_process_payment(email, amount):
    try:
        cursor.execute('''INSERT INTO members 
                        (name, institution, cluster, safaricom_number)
                        VALUES (?, ?, ?, ?)''',
                        ("Fahima Abdalla", "Timboni", "Cluster", "0727284993"))
        conn.commit()    
        client_secret = process_payment(amount, email)
        if client_secret:
            print("Payment processed successfully.")
        else:
            print("Failed to process payment.")
    except Exception as e:
        print(f"Failed to register member and process payment: {e}")
conn.close()

 
                            